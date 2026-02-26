import json
import os
import re
import google.generativeai as genai
from pydantic import BaseModel, Field
from typing import List, Optional
from dotenv import load_dotenv

load_dotenv()

# Ensure GEMINI_API_KEY is available in the environment variables
genai.configure(api_key=os.environ.get("GEMINI_API_KEY", ""))

# ---------------------------------------------------------------------------
# Input Schema definition (Chambers)
# ---------------------------------------------------------------------------
class ContactPerson(BaseModel):
    Name: Optional[str]
    Email: Optional[str]
    Telephone: Optional[str]

class DepartmentHead(BaseModel):
    Name: Optional[str]
    Email: Optional[str]
    Telephone: Optional[str]

class HireDeparture(BaseModel):
    Name: Optional[str]
    JoinedDeparted: Optional[str]
    JoinedFromDestination: Optional[str]

class Lawyer(BaseModel):
    Name: Optional[str]
    CommentsWebLink: Optional[str]
    PartnerYN: Optional[str]
    RankedYN: Optional[str]

class PublishableClient(BaseModel):
    Name: Optional[str]
    NewClient: Optional[str]

class Matter(BaseModel):
    ClientName: Optional[str]
    SummaryRole: Optional[str]
    Value: Optional[str]
    CrossBorderJurisdictions: Optional[str]
    LeadPartner: Optional[str]
    OtherTeamMembers: Optional[str]
    OtherFirmsAdvising: Optional[str]
    DateCompletion: Optional[str]

class ChambersInput(BaseModel):
    A1_FirmName: Optional[str]
    A2_PracticeArea: Optional[str]
    A3_Location: Optional[str]
    A4_ContactPerson: Optional[List[ContactPerson]]
    B1_DepartmentName: Optional[str]
    B2_PartnersCount: Optional[str]
    B3_OtherLawyersCount: Optional[str]
    B7_DepartmentHeads: Optional[List[DepartmentHead]]
    B8_HiresDepartures: Optional[List[HireDeparture]]
    B9_Lawyers: Optional[List[Lawyer]]
    B10_DepartmentKnownFor: Optional[str]
    D0_PublishableClients: Optional[List[PublishableClient]]
    D_PublishableMatters: Optional[List[Matter]]
    E0_ConfidentialClients: Optional[List[PublishableClient]]
    E_ConfidentialMatters: Optional[List[Matter]]

# ---------------------------------------------------------------------------
# Output Schema definition (Legal 500)
# ---------------------------------------------------------------------------
class L500ContactDetails(BaseModel):
    Name: Optional[str] = Field(None)
    JobTitle: Optional[str] = Field(None)
    Email: Optional[str] = Field(None)
    Phone: Optional[str] = Field(None)

class L500HeadsOfTeam(BaseModel):
    Name: Optional[str] = Field(None)
    Location: Optional[str] = Field(None)

class L500Client(BaseModel):
    Name: Optional[str] = Field(None)
    NewClient: Optional[str] = Field(None)

class L500Partner(BaseModel):
    Name: Optional[str] = Field(None)
    Location: Optional[str] = Field(None)
    RankedPrevious: Optional[str] = Field(None)
    SupportingInfo: Optional[str] = Field(None)

class L500HireDeparture(BaseModel):
    Name: Optional[str] = Field(None)
    Position: Optional[str] = Field(None)
    Status: Optional[str] = Field(None)
    Firm: Optional[str] = Field(None)
    Date: Optional[str] = Field(None)

class L500TeamMember(BaseModel):
    Name: Optional[str] = Field(None)
    Office: Optional[str] = Field(None)
    PracticeArea: Optional[str] = Field(None)

class L500OtherFirm(BaseModel):
    FirmName: Optional[str] = Field(None)
    RoleDetails: Optional[str] = Field(None)
    Advising: Optional[str] = Field(None)

class L500Matter(BaseModel):
    MatterType: Optional[str] = Field(None)
    ClientName: Optional[str] = Field(None)
    IndustrySector: Optional[str] = Field(None)
    MatterDescription: Optional[str] = Field(None)
    DealValue: Optional[str] = Field(None)
    CrossBorder: Optional[str] = Field(None)
    Jurisdictions: Optional[str] = Field(None)
    LeadPartners: Optional[List[L500TeamMember]] = Field(None)
    OtherKeyMembers: Optional[List[L500TeamMember]] = Field(None)
    OtherFirms: Optional[List[L500OtherFirm]] = Field(None)
    StartDate: Optional[str] = Field(None)
    EndDate: Optional[str] = Field(None)

class Legal500Output(BaseModel):
    FirmName: Optional[str] = Field(None)
    CountryLocation: Optional[str] = Field(None)
    PracticeArea: Optional[str] = Field(None)
    ContactDetails: Optional[List[L500ContactDetails]] = Field(None)
    TeamDepartmentName: Optional[str] = Field(None)
    HeadsOfTeam: Optional[List[L500HeadsOfTeam]] = Field(None)
    NumberPartners: Optional[str] = Field(None)
    NumberNonPartners: Optional[str] = Field(None)
    PracticeSetsApart: Optional[str] = Field(None)
    InitiativesInnovation: Optional[str] = Field(None)
    PublishableClients: Optional[List[L500Client]] = Field(None)
    NonPublishableClients: Optional[List[L500Client]] = Field(None)
    LeadingPartners: Optional[List[L500Partner]] = Field(None)
    NextGenerationPartners: Optional[List[L500Partner]] = Field(None)
    LeadingAssociates: Optional[List[L500Partner]] = Field(None)
    HiresDepartures: Optional[List[L500HireDeparture]] = Field(None)
    DetailedWorkHighlights: Optional[List[L500Matter]] = Field(None)

# ---------------------------------------------------------------------------
# Core Logic
# ---------------------------------------------------------------------------

def append_to_jsonl(record: dict, file_path: str = "records.jsonl"):
    """
    Appends a structured JSON dictionary to a JSONL file.
    """
    with open(file_path, 'a', encoding='utf-8') as f:
        f.write(json.dumps(record, ensure_ascii=False) + '\n')

def extract_count(text: str) -> str:
    """Helper to extract leading integer from strings like '12 | 58% Male'."""
    if not text:
        return ""
    match = re.search(r'^\s*(\d+)', str(text))
    return match.group(1) if match else str(text)

def parse_team_members(text: str) -> List[dict]:
    """Parses 'Name (Title), Name (Title)' into dicts."""
    if not text: return []
    members = []
    # Split by commas that aren't inside parentheses
    parts = [p.strip() for p in re.split(r',\s*(?![^()]*\))', str(text))]
    for part in parts:
        if not part: continue
        # Keep the title as part of the name
        members.append({
            "Name": part,
            "Office": "",
            "PracticeArea": ""
        })
    return members

def parse_other_firms(text: str) -> List[dict]:
    """Parses 'Skadden (advising DataCloud); ABNR (Indonesia)' into dicts."""
    if not text: return []
    firms = []
    parts = [p.strip() for p in str(text).split(';')]
    for part in parts:
        if not part: continue
        match = re.match(r'^(.*?)(?:\s*\((.*?)\))?$', part)
        if match:
            firm_name = match.group(1).strip()
            detail = match.group(2).strip() if match.group(2) else ""
            
            role_details = ""
            advising = ""
            
            if "advising" in detail.lower():
                advising = detail
            else:
                role_details = detail

            firms.append({
                "FirmName": firm_name,
                "RoleDetails": role_details,
                "Advising": advising
            })
    return firms

def process_matters(matters: List[dict], matter_type: str) -> List[dict]:
    output = []
    if not matters: return output
    for m in matters:
        out_m = {
            "MatterType": matter_type,
            "ClientName": m.get("ClientName", ""),
            "IndustrySector": "",
            "MatterDescription": m.get("SummaryRole", ""),
            "DealValue": m.get("Value", ""),
            "CrossBorder": "",
            "Jurisdictions": m.get("CrossBorderJurisdictions", ""),
            "StartDate": "",
            "EndDate": m.get("DateCompletion", ""),
            "LeadPartners": [{"Name": m.get("LeadPartner", ""), "Office": "", "PracticeArea": ""}] if m.get("LeadPartner") else [],
            "OtherKeyMembers": parse_team_members(m.get("OtherTeamMembers", "")),
            "OtherFirms": parse_other_firms(m.get("OtherFirmsAdvising", ""))
        }
        output.append(out_m)
    return output

def transform_chambers_to_legal500(chambers_data: dict) -> dict:
    """
    Transforms the ChambersInput dict into the Legal500Output dict shape
    using specific parsing and splitting rules.
    """
    input_data = chambers_data
    output = {}

    # 1. Direct Text Fields
    output["FirmName"] = input_data.get("A1_FirmName", "")
    output["PracticeArea"] = input_data.get("A2_PracticeArea", "")
    output["CountryLocation"] = input_data.get("A3_Location", "")
    
    # Contact Details
    output["ContactDetails"] = []
    for contact in input_data.get("A4_ContactPerson", []):
        output["ContactDetails"].append({
            "Name": contact.get("Name", ""),
            "JobTitle": "",
            "Email": contact.get("Email", ""),
            "Phone": contact.get("Telephone", "")
        })

    output["TeamDepartmentName"] = input_data.get("B1_DepartmentName", "")
    
    # Department Heads
    output["HeadsOfTeam"] = []
    for head in input_data.get("B7_DepartmentHeads", []):
        output["HeadsOfTeam"].append({
            "Name": head.get("Name", ""),
            "Location": ""
        })
        
    output["PracticeSetsApart"] = input_data.get("B10_DepartmentKnownFor", "")
    output["InitiativesInnovation"] = ""

    # 2. Number Extraction
    output["NumberPartners"] = extract_count(input_data.get("B2_PartnersCount", ""))
    output["NumberNonPartners"] = extract_count(input_data.get("B3_OtherLawyersCount", ""))

    # 3. Client Arrays
    output["PublishableClients"] = input_data.get("D0_PublishableClients", [])
    output["NonPublishableClients"] = input_data.get("E0_ConfidentialClients", [])

    # 4. Lawyer Array Splitting
    output["LeadingPartners"] = []
    output["NextGenerationPartners"] = []
    output["LeadingAssociates"] = []
    
    for lawyer in input_data.get("B9_Lawyers", []):
        partner_yn = str(lawyer.get("PartnerYN", "")).strip().upper()
        ranked_yn = str(lawyer.get("RankedYN", "")).strip().upper()
        
        info = {
            "Name": lawyer.get("Name", ""),
            "Location": "",
            "SupportingInfo": lawyer.get("CommentsWebLink", "")
        }
        
        if partner_yn == "Y" and ranked_yn == "Y":
            info["RankedPrevious"] = "Y"
            output["LeadingPartners"].append(info)
        elif partner_yn == "Y" and ranked_yn == "N":
            output["NextGenerationPartners"].append(info)
        elif partner_yn == "N":
            output["LeadingAssociates"].append(info)

    # 5. Hires and Departures Array
    output["HiresDepartures"] = []
    for hire in input_data.get("B8_HiresDepartures", []):
        output["HiresDepartures"].append({
            "Name": hire.get("Name", ""),
            "Position": "",
            "Status": hire.get("JoinedDeparted", ""),
            "Firm": hire.get("JoinedFromDestination", ""),
            "Date": ""
        })

    # 6. Matter Arrays
    matters = []
    matters.extend(process_matters(input_data.get("D_PublishableMatters", []), "Publishable matter"))
    matters.extend(process_matters(input_data.get("E_ConfidentialMatters", []), "Non-publishable matter"))
    output["DetailedWorkHighlights"] = matters

    return output

def copy_table_after(original_table):
    import copy
    from docx.oxml import OxmlElement
    from docx.table import Table
    new_tbl = copy.deepcopy(original_table._tbl)
    p = OxmlElement('w:p')
    original_table._tbl.addnext(p)
    p.addnext(new_tbl)
    return Table(new_tbl, original_table._parent)

def insert_row_after(table, base_row_idx):
    import copy
    from docx.table import _Row
    base_tr = table.rows[base_row_idx]._tr
    new_tr = copy.deepcopy(base_tr)
    base_tr.addnext(new_tr)
    # The new row is base_row_idx + 1
    return _Row(new_tr, table)

def fill_legal500_docx(data: dict, template_path: str, save_path: str):
    """
    Uses python-docx to fill an 'Output template (blank).docx' file
    using precise table indexes and dynamic row insertion.
    """
    from docx import Document
    doc = Document(template_path)
    
    # 0: Firm Name (Table 0)
    if len(doc.tables) > 0:
        doc.tables[0].cell(0,0).text = data.get("FirmName", "") or ""
        
    # Country & Practice Area TextBoxes
    # By inspecting the template, the Country textbox is in Paragraph 5, Practice Area in Paragraph 10
    if len(doc.paragraphs) > 10:
        country_p = doc.paragraphs[5]._element.xpath('.//w:txbxContent//w:p')
        if country_p:
            from docx.text.paragraph import Paragraph
            Paragraph(country_p[0], doc.paragraphs[5]).text = data.get("CountryLocation", "") or ""
            
        practice_p = doc.paragraphs[10]._element.xpath('.//w:txbxContent//w:p')
        if practice_p:
            from docx.text.paragraph import Paragraph
            Paragraph(practice_p[0], doc.paragraphs[10]).text = data.get("PracticeArea", "") or ""
        
    # 1: Contact Details (Row 1 to 3)
    contacts = data.get("ContactDetails", [])
    for i, c in enumerate(contacts[:3]):
        doc.tables[1].cell(1+i, 0).text = c.get('Name', '') or ""
        doc.tables[1].cell(1+i, 1).text = c.get('JobTitle', '') or ""
        doc.tables[1].cell(1+i, 2).text = c.get('Email', '') or ""
        doc.tables[1].cell(1+i, 3).text = c.get('Phone', '') or ""
        
    # 2: Team Department Name
    doc.tables[2].cell(0,0).text = data.get("TeamDepartmentName", "") or ""
        
    # 3: Head of Team
    heads = data.get("HeadsOfTeam", [])
    for i, h in enumerate(heads[:3]):
        doc.tables[3].cell(1+i, 0).text = h.get("Name", "") or ""
        doc.tables[3].cell(1+i, 1).text = h.get("Location", "") or ""
        
    # 4: Number of Partners / Non-Partners
    doc.tables[4].cell(1, 1).text = str(data.get("NumberPartners", ""))
    doc.tables[4].cell(1, 4).text = str(data.get("NumberNonPartners", ""))
    
    # 5 & 6: What sets your practice apart & Initiatives
    doc.tables[5].cell(0,0).text = data.get("PracticeSetsApart", "") or ""
    doc.tables[6].cell(0,0).text = data.get("InitiativesInnovation", "") or ""
    
    # 8 & 9: Clients
    pub_clients = data.get("PublishableClients", [])
    for i, c in enumerate(pub_clients[:10]):
        doc.tables[8].cell(1+i, 0).text = c.get("Name", "") or ""
        doc.tables[8].cell(1+i, 1).text = c.get("NewClient", "") or ""
        
    non_pub_clients = data.get("NonPublishableClients", [])
    for i, c in enumerate(non_pub_clients[:10]):
        doc.tables[9].cell(1+i, 0).text = c.get("Name", "") or ""
        doc.tables[9].cell(1+i, 1).text = c.get("NewClient", "") or ""
        
    # 11, 12, 13: Leading Partners
    leading_pts = data.get("LeadingPartners", [])
    for i, p in enumerate(leading_pts[:3]):
        t = doc.tables[11 + i]
        t.cell(2,0).text = p.get("Name", "") or ""
        t.cell(2,1).text = p.get("Location", "") or ""
        t.cell(2,2).text = p.get("RankedPrevious", "") or ""
        t.cell(4,0).text = p.get("SupportingInfo", "") or ""
        
    # 15, 16: Next Generation Partners
    next_gen_pts = data.get("NextGenerationPartners", [])
    for i, p in enumerate(next_gen_pts[:2]):
        t = doc.tables[15 + i]
        t.cell(2,0).text = p.get("Name", "") or ""
        t.cell(2,1).text = p.get("Location", "") or ""
        t.cell(2,2).text = p.get("RankedPrevious", "") or ""
        t.cell(4,0).text = p.get("SupportingInfo", "") or ""
        
    # 18, 19: Leading Associates
    lead_assoc = data.get("LeadingAssociates", [])
    for i, p in enumerate(lead_assoc[:2]):
        t = doc.tables[18 + i]
        t.cell(2,0).text = p.get("Name", "") or ""
        t.cell(2,1).text = p.get("Location", "") or ""
        t.cell(2,2).text = p.get("RankedPrevious", "") or ""
        t.cell(4,0).text = p.get("SupportingInfo", "") or ""

    def map_matter(t, highlight):
        # Top info is static
        t.cell(2,0).text = highlight.get("ClientName", "") or ""
        t.cell(2,2).text = highlight.get("IndustrySector", "") or ""
        t.cell(4,0).text = highlight.get("MatterDescription", "") or ""
        t.cell(5,1).text = highlight.get("DealValue", "") or ""
        juris = highlight.get("Jurisdictions", "") or ""
        cross = highlight.get("CrossBorder", "") or ""
        t.cell(8,0).text = f"{cross} - {juris}".strip('- ')
        
        # We need a robust cursor mapping to find labels, because inserting rows shifts indices.
        def find_row_by_label(table, label):
            for idx, r in enumerate(table.rows):
                if r.cells and label.lower() in r.cells[0].text.lower():
                    return idx
            return -1
            
        # -- LEAD PARTNERS --
        lead_idx = find_row_by_label(t, "Lead partner(s)")
        if lead_idx != -1:
            data_start = lead_idx + 2 # Skip label and headers
            partners = highlight.get("LeadPartners", [])
            # Template has 2 blank rows (data_start and data_start+1)
            for i in range(len(partners) - 2):
                insert_row_after(t, data_start)
            
            # Now fill them
            for i, p in enumerate(partners):
                row = t.rows[data_start + i]
                row.cells[0].text = p.get('Name', '') or ""
                row.cells[2].text = p.get('Office', '') or ""
                row.cells[4].text = p.get('PracticeArea', '') or ""
                
        # -- OTHER KEY TEAM MEMBERS --
        team_idx = find_row_by_label(t, "Other key team members")
        if team_idx != -1:
            data_start = team_idx + 2
            members = highlight.get("OtherKeyMembers", [])
            for i in range(len(members) - 2):
                insert_row_after(t, data_start)
                
            for i, m in enumerate(members):
                row = t.rows[data_start + i]
                row.cells[0].text = m.get('Name', '') or ""
                row.cells[2].text = m.get('Office', '') or ""
                row.cells[4].text = m.get('PracticeArea', '') or ""

        # -- OTHER FIRMS --
        firms_idx = find_row_by_label(t, "Other firms advising on the matter")
        if firms_idx != -1:
            data_start = firms_idx + 2
            firms = highlight.get("OtherFirms", [])
            # Template only has 1 blank row for firms
            for i in range(len(firms) - 1):
                insert_row_after(t, data_start)
                
            for i, f in enumerate(firms):
                row = t.rows[data_start + i]
                row.cells[0].text = f.get('FirmName', '') or ""
                row.cells[2].text = f.get('RoleDetails', '') or ""
                row.cells[4].text = f.get('Advising', '') or ""
                
        # -- START/END DATE --
        date_idx = find_row_by_label(t, "Start date")
        if date_idx != -1:
            data_start = date_idx + 1
            t.rows[data_start].cells[0].text = highlight.get("StartDate", "") or ""
            t.rows[data_start].cells[3].text = highlight.get("EndDate", "") or ""


    # 23+: Work Highlights
    highlights = data.get("DetailedWorkHighlights", [])
    
    # Split matters by Type
    pub_matters = [m for m in highlights if "non-publishable" not in m.get("MatterType", "").lower()]
    nonpub_matters = [m for m in highlights if "non-publishable" in m.get("MatterType", "").lower()]
    
    # Base tables inside the clean document
    base_pub_tbl = doc.tables[23]
    base_nonpub_tbl = doc.tables[24]
    
    # Process Publishable Matters
    for i, m in enumerate(pub_matters):
        if i == 0:
            active_tbl = base_pub_tbl
        else:
            # Clone new table after the previous active one
            active_tbl = copy_table_after(active_tbl)
        try:
            map_matter(active_tbl, m)
            active_tbl.cell(0,0).text = f"Publishable matter {i+1}"
        except Exception as e:
            print(f"Skipping publishable highlight {i} mapping due to table mismatch: {e}")
            
    # Process Non-publishable Matters
    for i, m in enumerate(nonpub_matters):
        if i == 0:
            active_tbl = base_nonpub_tbl
        else:
            active_tbl = copy_table_after(active_tbl)
        try:
            map_matter(active_tbl, m)
            active_tbl.cell(0,0).text = f"Non-publishable matter {i+1}"
        except Exception as e:
            print(f"Skipping non-publishable highlight {i} mapping due to table mismatch: {e}")
                
    doc.save(save_path)

def extract_chambers_data(input_data: str) -> dict:
    """
    Calls the Gemini API using structured outputs to extract the text into ChambersInput.
    """
    model = genai.GenerativeModel("gemini-2.5-flash")
    
    prompt = f"""
    You are an AI tasked with translating raw text from a document format into 
    the provided structured JSON schema. Extract all the information cleanly.
    
    Use the exact text if possible. Handle missing text gracefully by returning NULL or empty string.
    
    You MUST return ONLY valid JSON matching this exact schema:
    {json.dumps(ChambersInput.model_json_schema(), indent=2)}
    
    Input Document Context:
    ---
    {input_data}
    ---
    """
    
    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                temperature=0.0
            ),
        )
        return json.loads(response.text)
    except Exception as e:
        print(f"Error during Gemini call: {e}")
        raise

def handler(event, context):
    """
    AWS Lambda handler function.
    """
    try:
        # Check if the body comes from an API Gateway Proxy event or direct invocation
        if 'body' in event:
            if isinstance(event['body'], str):
                input_data = event['body']
            else:
                input_data = json.dumps(event['body'])
        else:
            input_data = json.dumps(event)

        if not input_data:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "No input data provided."})
            }

        # Perform extraction
        structured_output = extract_legal500_data(input_data)
        
        return {
            "statusCode": 200,
            "headers": {
                "Content-Type": "application/json"
            },
            "body": json.dumps(structured_output)
        }
        
    except Exception as e:
        print(f"Error processing request: {e}")
        return {
            "statusCode": 500,
            "body": json.dumps({"error": str(e)})
        }
