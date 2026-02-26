import streamlit as st
import json
import os
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv
from app import extract_chambers_data, transform_chambers_to_legal500, append_to_jsonl, fill_legal500_docx

load_dotenv()

st.set_page_config(page_title="Form Translator", layout="wide")

st.title("Chambers to Legal 500 Translator")
st.markdown("Automated end-to-end pipeline: Upload `.docx`, extract to `.jsonl`, and generate `Legal 500 Output.docx`.")

# Check for API Key
if "GEMINI_API_KEY" in os.environ:
    st.sidebar.success("API Key detected in environment variables.")

# ---------------------------------------------------------------------------
# Left Column: Uploads
# ---------------------------------------------------------------------------
st.header("Step 1: Input (Chambers)")
uploaded_files = st.file_uploader("Upload Chambers Document(s)", type=["docx"], accept_multiple_files=True)

if "results" not in st.session_state:
    st.session_state["results"] = []

if st.button("Extract Data", type="primary"):
    st.session_state["results"] = [] # Reset on rerun
    
    if uploaded_files:
        with st.spinner("Extracting fields with Gemini 2.5 Flash..."):
            for i, uploaded_file in enumerate(uploaded_files):
                st.write(f"Processing: `{uploaded_file.name}`")
                
                # Extract text
                extracted_text = ""
                try:
                    doc = Document(uploaded_file)
                    for para in doc.paragraphs:
                        extracted_text += para.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                extracted_text += cell.text + " | "
                            extracted_text += "\n"
                except Exception as e:
                    st.error(f"Error reading docx: {e}")
                    continue
                    
                # Call Gemini & Transform
                try:
                    chambers_result = extract_chambers_data(extracted_text)
                    legal500_result = transform_chambers_to_legal500(chambers_result)
                    st.session_state["results"].append({
                        "filename": uploaded_file.name,
                        "data": legal500_result
                    })
                    st.success(f"Successfully extracted {uploaded_file.name}")
                except Exception as model_err:
                    st.error(f"Translation logic error on {uploaded_file.name}:\n{model_err}")
    else:
        st.warning("Please upload at least one .docx file.")

# ---------------------------------------------------------------------------
# Middle/Right section: Generation
# ---------------------------------------------------------------------------
if st.session_state["results"]:
    st.header("Step 2: Review and Save")
    st.write(f"Extracted {len(st.session_state['results'])} documents.")
    
    for res in st.session_state["results"]:
        with st.expander(res["filename"]):
            st.json(res["data"])

    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Save to JSONL (`records.jsonl`)"):
            for res in st.session_state["results"]:
                append_to_jsonl(res["data"])
            st.success("Saved all extracted data to records.jsonl!")
            
    with col2:
        if st.button("Generate Legal 500 Output DOCX"):
            # Ensure the blank template exists
            template_path = "data/Output template (blank).docx"
            if not os.path.exists(template_path):
                st.error(f"Template not found at `{template_path}`. Please make sure the 'data' folder has the blank template.")
            else:
                for res in st.session_state["results"]:
                    # Create a safe output filename based on input
                    out_name = f"Output_Generated_{res['filename']}"
                    
                    try:
                        fill_legal500_docx(res["data"], template_path, out_name)
                        st.success(f"Generated successfully: `{out_name}`")
                        
                        # Provide Download Button for the generated DOCX
                        with open(out_name, "rb") as file:
                            st.download_button(
                                label=f"Download {out_name}",
                                data=file,
                                file_name=out_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Failed to fill Document `{res['filename']}`: {e}")
